#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
读取开题报告，了解表格结构
"""

from docx import Document

def read_report_structure(file_path):
    """读取报告结构，查看表格"""
    doc = Document(file_path)
    
    print(f"文档包含 {len(doc.paragraphs)} 个段落")
    print(f"文档包含 {len(doc.tables)} 个表格")
    
    # 遍历所有表格
    for i, table in enumerate(doc.tables):
        print(f"\n表格 {i+1}:")
        print(f"  行数: {len(table.rows)}")
        print(f"  列数: {len(table.columns)}")
        
        # 打印表格内容
        for j, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            print(f"  行 {j+1}: {cells}")

if __name__ == "__main__":
    file_path = "d:\\code\\idapromcp_333-main\\王晨俊毕业论文（设计）开题报告 .docx"
    read_report_structure(file_path)
